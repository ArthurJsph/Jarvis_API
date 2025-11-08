"""
Manipulação de Documentos Word
Jarvis CLI - Módulo de Word
"""

import os
from pathlib import Path
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import config
from rich.console import Console

console = Console()


class WordManager:
    """Gerenciador de documentos Word"""
    
    def __init__(self):
        self.encoding = config.DEFAULT_ENCODING
        self.default_author = config.DEFAULT_AUTHOR
        self.default_organization = config.DEFAULT_ORGANIZATION
    
    def create_document(self, filename: str, title: str = "", content: str = "") -> bool:
        """
        Cria um novo documento Word
        
        Args:
            filename: Nome do arquivo
            title: Título do documento
            content: Conteúdo inicial
        """
        try:
            doc = Document()
            
            # Adiciona título se fornecido
            if title:
                title_paragraph = doc.add_heading(title, level=1)
                title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Adiciona conteúdo se fornecido
            if content:
                doc.add_paragraph(content)
            
            # Adiciona informações do autor
            core_props = doc.core_properties
            core_props.author = self.default_author
            if self.default_organization:
                core_props.category = self.default_organization
            core_props.comments = f"Documento criado pelo {config.APP_NAME}"
            
            # Salva o documento
            filepath = Path(config.DOCUMENTS_DIR) / filename
            if not filepath.suffix:
                filepath = filepath.with_suffix('.docx')
            
            doc.save(filepath)
            
            console.print(f"✅ Documento Word criado: {filepath}", style="green")
            return True
            
        except Exception as e:
            console.print(f"❌ Erro ao criar documento Word: {str(e)}", style="red")
            return False
    
    def read_document(self, filename: str) -> Optional[str]:
        """
        Lê o conteúdo de um documento Word
        
        Args:
            filename: Nome do arquivo
        """
        try:
            filepath = Path(filename)
            if not filepath.exists():
                filepath = Path(config.DOCUMENTS_DIR) / filename
                
            if not filepath.exists():
                console.print(f"❌ Arquivo não encontrado: {filename}", style="red")
                return None
            
            doc = Document(filepath)
            
            # Extrai todo o texto
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            
            content = '\n'.join(full_text)
            
            console.print(f"📄 Documento lido: {filepath}", style="green")
            console.print(f"Caracteres: {len(content)}", style="cyan")
            
            return content
            
        except Exception as e:
            console.print(f"❌ Erro ao ler documento Word: {str(e)}", style="red")
            return None
    
    def extract_to_text(self, filename: str, output_file: Optional[str] = None) -> bool:
        """
        Extrai texto de um documento Word para arquivo .txt
        
        Args:
            filename: Nome do arquivo Word
            output_file: Nome do arquivo de saída (opcional)
        """
        try:
            content = self.read_document(filename)
            if not content:
                return False
            
            if not output_file:
                input_path = Path(filename)
                if not input_path.exists():
                    input_path = Path(config.DOCUMENTS_DIR) / filename
                output_file = input_path.stem + "_extracted.txt"
            
            output_path = Path(config.DOCUMENTS_DIR) / output_file
            
            with open(output_path, 'w', encoding=self.encoding) as f:
                f.write(content)
            
            console.print(f"✅ Texto extraído para: {output_path}", style="green")
            return True
            
        except Exception as e:
            console.print(f"❌ Erro ao extrair texto: {str(e)}", style="red")
            return False
    
    def find_and_replace(self, filename: str, find_text: str, replace_text: str, 
                        output_file: Optional[str] = None) -> bool:
        """
        Busca e substitui texto em um documento Word
        
        Args:
            filename: Nome do arquivo
            find_text: Texto a ser encontrado
            replace_text: Texto de substituição
            output_file: Arquivo de saída (opcional, senão substitui o original)
        """
        try:
            filepath = Path(filename)
            if not filepath.exists():
                filepath = Path(config.DOCUMENTS_DIR) / filename
                
            if not filepath.exists():
                console.print(f"❌ Arquivo não encontrado: {filename}", style="red")
                return False
            
            doc = Document(filepath)
            replacements = 0
            
            # Substitui em parágrafos
            for paragraph in doc.paragraphs:
                if find_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(find_text, replace_text)
                    replacements += 1
            
            # Substitui em tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if find_text in cell.text:
                            cell.text = cell.text.replace(find_text, replace_text)
                            replacements += 1
            
            # Salva o documento
            if output_file:
                output_path = Path(config.DOCUMENTS_DIR) / output_file
            else:
                output_path = filepath
            
            doc.save(output_path)
            
            console.print(f"✅ {replacements} substituições feitas em: {output_path}", style="green")
            return True
            
        except Exception as e:
            console.print(f"❌ Erro na busca e substituição: {str(e)}", style="red")
            return False
    
    def bulk_find_replace(self, directory: str, find_text: str, replace_text: str) -> int:
        """
        Busca e substitui texto em vários arquivos Word
        
        Args:
            directory: Diretório com os arquivos
            find_text: Texto a ser encontrado
            replace_text: Texto de substituição
        """
        try:
            dir_path = Path(directory)
            if not dir_path.exists():
                dir_path = Path(config.DOCUMENTS_DIR) / directory
            
            docx_files = list(dir_path.glob("*.docx"))
            
            if not docx_files:
                console.print(f"❌ Nenhum arquivo .docx encontrado em: {dir_path}", style="red")
                return 0
            
            total_files = 0
            for docx_file in docx_files:
                if self.find_and_replace(str(docx_file), find_text, replace_text):
                    total_files += 1
            
            console.print(f"✅ Processados {total_files} arquivos", style="green")
            return total_files
            
        except Exception as e:
            console.print(f"❌ Erro no processamento em lote: {str(e)}", style="red")
            return 0
    
    def create_from_template(self, template_data: Dict[str, Any], 
                           template_file: Optional[str] = None, 
                           output_file: str = "document_from_template.docx") -> bool:
        """
        Cria documento a partir de template com dados dinâmicos
        
        Args:
            template_data: Dados para preencher o template
            template_file: Arquivo de template (opcional)
            output_file: Arquivo de saída
        """
        try:
            if template_file and Path(template_file).exists():
                doc = Document(template_file)
            else:
                # Cria template básico
                doc = Document()
                
                # Template padrão
                doc.add_heading(template_data.get('title', 'Documento'), level=1)
                
                if 'author' in template_data:
                    doc.add_paragraph(f"Autor: {template_data['author']}")
                
                if 'date' in template_data:
                    doc.add_paragraph(f"Data: {template_data['date']}")
                
                doc.add_paragraph()  # Linha em branco
                
                if 'content' in template_data:
                    doc.add_paragraph(template_data['content'])
                
                if 'sections' in template_data:
                    for section in template_data['sections']:
                        doc.add_heading(section.get('title', 'Seção'), level=2)
                        doc.add_paragraph(section.get('content', ''))
            
            # Substitui placeholders se existirem
            for key, value in template_data.items():
                placeholder = f"{{{key}}}"
                for paragraph in doc.paragraphs:
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, str(value))
            
            # Salva o documento
            output_path = Path(config.DOCUMENTS_DIR) / output_file
            doc.save(output_path)
            
            console.print(f"✅ Documento criado a partir de template: {output_path}", style="green")
            return True
            
        except Exception as e:
            console.print(f"❌ Erro ao criar documento do template: {str(e)}", style="red")
            return False
    
    def add_table_to_document(self, filename: str, table_data: List[List[str]], 
                             headers: Optional[List[str]] = None) -> bool:
        """
        Adiciona uma tabela a um documento existente
        
        Args:
            filename: Nome do arquivo
            table_data: Dados da tabela
            headers: Cabeçalhos da tabela
        """
        try:
            filepath = Path(filename)
            if not filepath.exists():
                filepath = Path(config.DOCUMENTS_DIR) / filename
            
            if filepath.exists():
                doc = Document(filepath)
            else:
                doc = Document()
            
            # Determina o número de colunas
            cols = len(headers) if headers else len(table_data[0]) if table_data else 1
            rows = len(table_data) + (1 if headers else 0)
            
            # Cria a tabela
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            
            # Adiciona cabeçalhos
            if headers:
                header_row = table.rows[0]
                for i, header in enumerate(headers):
                    header_row.cells[i].text = header
                    # Formata cabeçalho
                    for paragraph in header_row.cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
            
            # Adiciona dados
            start_row = 1 if headers else 0
            for i, row_data in enumerate(table_data):
                row = table.rows[start_row + i]
                for j, cell_data in enumerate(row_data):
                    if j < len(row.cells):
                        row.cells[j].text = str(cell_data)
            
            # Salva o documento
            doc.save(filepath)
            
            console.print(f"✅ Tabela adicionada ao documento: {filepath}", style="green")
            return True
            
        except Exception as e:
            console.print(f"❌ Erro ao adicionar tabela: {str(e)}", style="red")
            return False
    
    def get_document_info(self, filename: str) -> Dict[str, Any]:
        """
        Obtém informações sobre um documento Word
        
        Args:
            filename: Nome do arquivo
        """
        try:
            filepath = Path(filename)
            if not filepath.exists():
                filepath = Path(config.DOCUMENTS_DIR) / filename
                
            if not filepath.exists():
                console.print(f"❌ Arquivo não encontrado: {filename}", style="red")
                return {}
            
            doc = Document(filepath)
            
            # Conta elementos
            paragraphs = len(doc.paragraphs)
            tables = len(doc.tables)
            
            # Conta palavras e caracteres
            full_text = '\n'.join([p.text for p in doc.paragraphs])
            words = len(full_text.split())
            characters = len(full_text)
            
            # Informações do arquivo
            file_size = filepath.stat().st_size
            
            info = {
                "filename": filepath.name,
                "size": f"{file_size / 1024:.2f} KB",
                "paragraphs": paragraphs,
                "tables": tables,
                "words": words,
                "characters": characters,
                "author": doc.core_properties.author or "Desconhecido",
                "created": doc.core_properties.created,
                "modified": doc.core_properties.modified
            }
            
            return info
            
        except Exception as e:
            console.print(f"❌ Erro ao obter info do documento: {str(e)}", style="red")
            return {}